import os
import re
import pdfplumber
import docx

def extract_text_from_pdf(file_path):
    """Extract all text from a PDF file using pdfplumber."""
    text = ""
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        print(f"Error reading PDF: {e}")
    return text

def extract_text_from_docx(file_path):
    """Extract all text from a DOCX file, including tables."""
    text_content = []
    try:
        doc = docx.Document(file_path)
        # Extract paragraph text
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text.strip())
        
        # Extract table text
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    # Prevent duplicating text if cells are merged
                    cell_text = cell.text.strip()
                    if cell_text and cell_text not in text_content:
                        text_content.append(cell_text)
                        
    except Exception as e:
        print(f"Error reading DOCX: {e}")
    return "\n".join(text_content)

def clean_text(text):
    """Normalize whitespace, remove special chars, and return cleaned text."""
    if not text:
        return ""
    # Replace non-breaking spaces and tabs with space
    text = text.replace('\xa0', ' ').replace('\t', ' ')
    # Normalize multiple newlines and spaces
    text = re.sub(r'\s+', ' ', text)
    return text.strip()

def detect_sections(text):
    """
    Detect the presence of standard resume sections.
    Returns a dictionary of section -> boolean.
    """
    sections = {
        "Education": False,
        "Skills": False,
        "Projects": False,
        "Experience": False,
        "Certifications": False
    }
    
    if not text:
        return sections

    # We will search using regular expressions for common headers
    # Headers are usually on their own line or at the start of a block
    lines = [line.strip().lower() for line in text.split('\n') if line.strip()]
    
    # Section keyword patterns
    patterns = {
        "Education": [r'\beducation\b', r'\bacademics?\b', r'\bacademic background\b', r'\bqualifications?\b', r'\bschooling\b'],
        "Skills": [r'\bskills?\b', r'\btechnical skills?\b', r'\btechnologies\b', r'\bcore competencies\b', r'\bexpertise\b', r'\btools\b'],
        "Projects": [r'\bprojects?\b', r'\bacacademic projects?\b', r'\bpersonal projects?\b', r'\bkey projects?\b'],
        "Experience": [r'\bexperience\b', r'\bwork experience\b', r'\bprofessional experience\b', r'\bemployment\b', r'\bwork history\b', r'\binternships?\b'],
        "Certifications": [r'\bcertifications?\b', r'\bcertificates?\b', r'\bcourses?\b', r'\blicenses?\b', r'\baccreditation\b']
    }

    # First check: Search by whole line matches (higher confidence headers)
    for section, pattern_list in patterns.items():
        for line in lines:
            for pattern in pattern_list:
                # Matches patterns like "Education", "Education & Certifications", etc.
                if re.search(r'^' + pattern, line) or re.search(pattern + r'$', line):
                    sections[section] = True
                    break
            if sections[section]:
                break
                
    # Second check fallback: general regex search in the entire text if not found
    for section, pattern_list in patterns.items():
        if not sections[section]:
            for pattern in pattern_list:
                if re.search(pattern, text.lower()):
                    sections[section] = True
                    break
                    
    return sections
