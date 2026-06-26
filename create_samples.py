import os
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
import docx

def create_pdf_resume(filepath):
    """Generate a sample AI/ML Engineer resume PDF using ReportLab."""
    doc = SimpleDocTemplate(
        filepath,
        pagesize=letter,
        leftMargin=54,
        rightMargin=54,
        topMargin=54,
        bottomMargin=54
    )
    
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle(
        'Name',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=20,
        leading=24,
        textColor=colors.HexColor("#1A1A24")
    )
    
    subtitle_style = ParagraphStyle(
        'Subtitle',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=10,
        leading=14,
        textColor=colors.HexColor("#555555")
    )
    
    h1_style = ParagraphStyle(
        'SectionHeader',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=13,
        leading=16,
        textColor=colors.HexColor("#4A154B"),
        spaceBefore=10,
        spaceAfter=4
    )
    
    body_style = ParagraphStyle(
        'Body',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9.5,
        leading=13.5,
        textColor=colors.HexColor("#333333"),
        spaceAfter=4
    )
    
    story = []
    
    # Header
    story.append(Paragraph("JOHN DOE", title_style))
    story.append(Paragraph("AI & Machine Learning Engineer | john.doe@email.com | +1-234-567-8900 | San Francisco, CA", subtitle_style))
    story.append(Spacer(1, 10))
    
    # Education
    story.append(Paragraph("Education", h1_style))
    story.append(Paragraph("<b>M.S. in Computer Science (Specialization in AI)</b> - Stanford University | 2022 - 2024", body_style))
    story.append(Paragraph("<b>B.S. in Computer Science</b> - California State University | 2018 - 2022", body_style))
    story.append(Spacer(1, 8))
    
    # Experience
    story.append(Paragraph("Experience", h1_style))
    story.append(Paragraph("<b>Machine Learning Engineer</b> - DeepMind Solutions | Jan 2023 - Present", body_style))
    story.append(Paragraph("- Designed and trained Deep Learning architectures for image classification using PyTorch.", body_style))
    story.append(Paragraph("- Handled Data Science pipelines cleaning over 10M rows of raw text with Pandas and NumPy.", body_style))
    story.append(Paragraph("- Deployed transformer-based NLP applications using Docker and AWS SageMaker.", body_style))
    story.append(Spacer(1, 8))
    
    # Projects
    story.append(Paragraph("Projects", h1_style))
    story.append(Paragraph("<b>Medical Image Segmentation App</b> | Python, PyTorch, NumPy", body_style))
    story.append(Paragraph("- Built a Deep Learning U-Net framework for cell boundary segmentation with 94% accuracy.", body_style))
    story.append(Paragraph("- Utilized Scikit-learn for baseline classifier comparison and training telemetry validation.", body_style))
    story.append(Spacer(1, 8))
    
    # Skills
    story.append(Paragraph("Skills", h1_style))
    story.append(Paragraph("<b>Programming:</b> Python, Java, SQL, C++", body_style))
    story.append(Paragraph("<b>AI/ML:</b> Machine Learning, Deep Learning, Natural Language Processing, NLP, TensorFlow, PyTorch, Scikit-learn, Pandas, NumPy", body_style))
    story.append(Paragraph("<b>Tools/Cloud:</b> Git, Docker, AWS, GCP, Linux", body_style))
    story.append(Spacer(1, 8))
    
    # Certifications
    story.append(Paragraph("Certifications", h1_style))
    story.append(Paragraph("- AWS Certified Machine Learning Specialty (2024)", body_style))
    story.append(Paragraph("- DeepLearning.AI Deep Learning Specialization (2023)", body_style))
    
    doc.build(story)

def create_docx_resume(filepath):
    """Generate a sample Full Stack Web Developer resume DOCX using python-docx."""
    doc = docx.Document()
    
    # Set spacing
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = docx.shared.Pt(10.5)
    
    # Name
    name = doc.add_paragraph()
    r = name.add_run("JANE SMITH")
    r.bold = True
    r.font.size = docx.shared.Pt(20)
    name.alignment = 0 # Left
    
    # Contact
    contact = doc.add_paragraph("Full Stack Web Developer | jane.smith@email.com | +1-987-654-3210 | New York, NY")
    contact.runs[0].font.size = docx.shared.Pt(9.5)
    contact.runs[0].font.color.rgb = docx.shared.RGBColor(100, 100, 100)
    
    # Education
    doc.add_heading("Education", level=1)
    p_edu = doc.add_paragraph()
    p_edu.add_run("B.Tech in Computer Science and Engineering\n").bold = True
    p_edu.add_run("University of Maryland | 2019 - 2023 | GPA: 3.8/4.0")
    
    # Experience
    doc.add_heading("Experience", level=1)
    p_exp = doc.add_paragraph()
    p_exp.add_run("Frontend Developer Intern\n").bold = True
    p_exp.add_run("ByteCraft Technologies | June 2022 - May 2023\n")
    p_exp.add_run("- Developed responsive dashboard interfaces using JavaScript, HTML, CSS, and React.\n")
    p_exp.add_run("- Collaborated with backend teams to integrate REST APIs built in Node.js and Express.\n")
    p_exp.add_run("- Conducted version control workflows utilizing Git and GitHub repository branching.")
    
    # Projects
    doc.add_heading("Projects", level=1)
    p_proj = doc.add_paragraph()
    p_proj.add_run("E-Commerce Web Portal\n").bold = True
    p_proj.add_run("- Programmed a dynamic online storefront utilizing React for stateful page UI.\n")
    p_proj.add_run("- Leveraged PostgreSQL database models and Docker container setups for development scaling.")
    
    # Skills
    doc.add_heading("Skills", level=1)
    p_skills = doc.add_paragraph()
    p_skills.add_run("Web Technologies: ").bold = True
    p_skills.add_run("HTML, CSS, JavaScript, TypeScript, React, Bootstrap, Tailwind\n")
    p_skills.add_run("Backend/Database: ").bold = True
    p_skills.add_run("Node.js, Express, Python, SQL, PostgreSQL, MongoDB\n")
    p_skills.add_run("DevOps/Cloud: ").bold = True
    p_skills.add_run("Git, Docker, AWS, Jenkins, Linux")
    
    # Certifications
    doc.add_heading("Certifications", level=1)
    p_cert = doc.add_paragraph()
    p_cert.add_run("- AWS Certified Solutions Architect Associate (2024)\n")
    p_cert.add_run("- Meta Front-End Developer Professional Certificate (2023)")
    
    doc.save(filepath)

def create_jd_txt(filepath, jd_type):
    """Generate sample job description text files."""
    if jd_type == "ai":
        content = """Position: AI & Machine Learning Engineer

Company: AI Technologies Inc.
Location: San Francisco, CA (Hybrid)

Responsibilities:
- Build, optimize, and train Machine Learning and Deep Learning algorithms.
- Build clean Data Science pipelines using Pandas, NumPy, and Scikit-learn.
- Train state-of-the-art Neural Networks using PyTorch or TensorFlow.
- Design Natural Language Processing (NLP) modules for semantic matching.
- Containerize deployment modules with Docker and deploy on AWS cloud infrastructures.

Requirements:
- Strong programming experience in Python, Java, and SQL.
- Solid background in Machine Learning, Deep Learning, and NLP architectures.
- Experience with Docker, Git version control, and AWS SageMaker.
- Master's or Bachelor's degree in Computer Science, Data Science, or related field.
"""
    else:
        content = """Position: Frontend Web Developer (React)

Company: WebCraft Systems
Location: New York, NY

Responsibilities:
- Build responsive, interactive web interfaces using React, JavaScript, HTML, and CSS.
- Optimize frontend components for speed and scalability.
- Work closely with backend teams to query database API layers.
- Maintain code repositories using Git and manage deployment configurations with Docker.

Requirements:
- Proficiency in HTML, CSS, JavaScript, and React.
- Understanding of backend integrations, Node.js, and SQL or PostgreSQL databases.
- Experience with Git, GitHub, Docker, and AWS deployments.
- Strong UI/UX aesthetics and attention to mobile-responsive structures.
"""
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(content)

if __name__ == "__main__":
    os.makedirs("samples", exist_ok=True)
    
    create_pdf_resume("samples/john_doe_ai_resume.pdf")
    create_docx_resume("samples/jane_smith_web_resume.docx")
    create_jd_txt("samples/ai_engineer_jd.txt", "ai")
    create_jd_txt("samples/frontend_developer_jd.txt", "web")
    
    print("Sample resumes and job descriptions created successfully in the 'samples/' directory!")
